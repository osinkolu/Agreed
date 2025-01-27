SNOWFLAKE_CONFIG = {
    "user": "xxxxxxxx",
    "password": "xxxxxxxx",
    "account": "xxxxxxxxx",
    "warehouse": "COMPUTE_WH",
    "database": "CC_QUICKSTART_CORTEX_SEARCH_DOCS",
    "schema": "CONTRACTS",
    "search_service": "CC_SEARCH_SERVICE_CS"
}

from flask import Flask, request, jsonify
import os
# import snowflake.connector
from snowflake.core import Root
from snowflake.snowpark import Session
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
from docx import Document
from werkzeug.utils import secure_filename

# Flask App Initialization
app = Flask(__name__)

# Snowflake Session Initialization
session = Session.builder.configs(SNOWFLAKE_CONFIG).create()

# Chunking Configuration
CHUNK_SIZE = 5000
CHUNK_OVERLAP = 256

# Allowed Extensions
ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    with open(file_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        return "\n".join(page.extract_text() for page in reader.pages)

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def save_chunks_to_snowflake(chunks, metadata):
    """Save document chunks and metadata to Snowflake."""
    for idx, chunk in enumerate(chunks):
        session.sql(
            "INSERT INTO DOCUMENT_CHUNKS (chunk, tag_name, filename, page_number, paragraph_number, username) VALUES (?, ?, ?, ?, ?, ?)",
            params=[chunk, metadata["tag_name"], metadata["filename"], metadata["page_number"], idx + 1, metadata["username"]]
        ).collect()

@app.route('/')
def hello_world():
    return("Welcome, please smile more 😜")

@app.route("/api/save-document", methods=["POST"])
def save_document():
    if "file" not in request.files:
        return jsonify({"error": "No file part in the request."}), 400
    
    if "username" not in request.form:
        return jsonify({"error": "No username in the request."}), 400    

    file = request.files["file"]
    tag_name = request.form.get("tag_name")
    username = request.form.get("username")

    if file.filename == "":
        return jsonify({"error": "No selected file."}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        # file_path = os.path.join("/tmp", filename)
        # file.save(file_path)

        upload_dir = "/tmp"
        # Ensure the upload directory exists
        os.makedirs(upload_dir, exist_ok=True)
        # Create the full file path
        file_path = os.path.join(upload_dir, filename)
        # Save the file
        file.save(file_path)

        # Extract text based on file type
        try:
            if filename.endswith(".pdf"):
                full_text = extract_text_from_pdf(file_path)
            elif filename.endswith(".docx"):
                full_text = extract_text_from_docx(file_path)
            else:
                return jsonify({"error": "Unsupported file format."}), 400

            # Chunk the document
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
            )
            chunks = splitter.split_text(full_text)

            # Save chunks to Snowflake
            metadata = {
                "tag_name": tag_name,
                "filename": filename,
                "page_number": 1,  # Assuming page is irrelevant for DOCX
                "username": username
            }
            save_chunks_to_snowflake(chunks, metadata)

            # Generate summary
            summary = session.sql(
                "SELECT SNOWFLAKE.CORTEX.SUMMARIZE(?) AS summary", params=[full_text]
            ).collect()[0]["SUMMARY"]

            return jsonify({"message": "Document saved successfully.", "summary": summary})
        finally:
            os.remove(file_path)

    return jsonify({"error": "File type not allowed."}), 400

@app.route("/api/context-retrieval", methods=["POST"])

def context_retrieval():
    data = request.json
    username = data.get("username")
    query = data.get("query")
    n = data.get("n", 5)

    if not query:
        return jsonify({"error": "Query is required."}), 400

    
    if "username" not in data:
        return jsonify({"error": "No username in the request."}), 400  
    
    try:
        # Initialize the Cortex Search Service
        root = Root(session)
        svc = root.databases[SNOWFLAKE_CONFIG["database"]].schemas[SNOWFLAKE_CONFIG["schema"]].cortex_search_services[SNOWFLAKE_CONFIG["search_service"]]

        # Build Filter Object
        filter_obj = {"@eq": {"username": username}} if username != "All" else None
        
        # Perform the search query
        response = svc.search(
            query=query,
            columns=["chunk", "tag_name", "filename", "page_number", "paragraph_number"],
            filter=filter_obj,
            limit=n
        )

        # Process results
        results = [
            {
                "chunk": result["chunk"],
                "tag_name": result["tag_name"],
                "filename": result["filename"],
                "page_number": result["page_number"],
                "paragraph_number": result["paragraph_number"]
            }
            for result in response.results
        ]

        return jsonify({"results": results})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/conflict-checker", methods=["POST"])
def conflict_checker():
    if "file" not in request.files:
        return jsonify({"error": "No file part in the request."}), 400
    
    if "username" not in request.form:
        return jsonify({"error": "No username in the request."}), 400  

    file = request.files["file"]
    username = request.form.get("username")

    if file.filename == "":
        return jsonify({"error": "No selected file."}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join("/tmp", filename)
        file.save(file_path)

        # Extract text from file
        try:
            if filename.endswith(".pdf"):
                full_text = extract_text_from_pdf(file_path)
            elif filename.endswith(".docx"):
                full_text = extract_text_from_docx(file_path)
            else:
                return jsonify({"error": "Unsupported file format."}), 400

            # Chunk the extracted text into smaller pieces
            splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
            chunks = splitter.split_text(full_text)

            # Store conflicts in a list
            conflicts = []

            # Process each chunk
            for chunk_idx, chunk in enumerate(chunks):
                # Use the existing context retrieval function to get similar content
                query_data = {
                    "query": chunk,
                    "n": 5,  # Retrieve top 5 similar chunks
                    "username":username
                }
                with app.test_request_context(json=query_data):
                    context_response = context_retrieval()
                    if context_response.status_code != 200:
                        return jsonify({"error": f"Error retrieving context for chunk {chunk_idx + 1}."}), 500

                    # Extract retrieved context from the response
                    retrieved_context = "\n".join(
                        [result["chunk"] for result in context_response.get_json()["results"]]
                    )

                # Generate the query prompt for conflict detection
                query_prompt = f"""
                You are an AI assistant. Use the CONTEXT below to check for conflicts with the NEW CONTRACT CHUNK.

                CONTEXT:
                {retrieved_context}

                NEW CONTRACT CHUNK:
                {chunk}

                QUESTION:
                Does the new contract chunk conflict with any of the agreements in the context? 
                If so, provide the details of the conflict.

                ANSWER:
                """

                # Generate the response using Snowflake Cortex
                try:
                    print("LFG BABY!!!!!!!!!!!")
                    result = session.sql(
                        "SELECT snowflake.cortex.complete(?, ?) AS response",
                        params=["mistral-large2", query_prompt]
                    ).collect()

                    answer = result[0]["RESPONSE"]

                    # If conflicts are found, add them to the list
                    if answer.strip():  # Check if the response contains meaningful conflict information
                        conflicts.append(f"Chunk {chunk_idx + 1}: {answer}")

                except Exception as e:
                    return jsonify({"error": f"Error processing chunk {chunk_idx + 1}: {str(e)}"}), 500

            # Summarize all conflicts
            full_conflicts_text = "\n".join(conflicts)
            if full_conflicts_text.strip():  # Only summarize if there are conflicts
                summary = session.sql(
                    "SELECT SNOWFLAKE.CORTEX.SUMMARIZE(?) AS summary",
                    params=[full_conflicts_text]
                ).collect()[0]["SUMMARY"]
            else:
                summary = "No conflicts detected across the document."

            # Return the conflict summary
            return jsonify({"conflict_summary": summary, "complete_conflict":full_conflicts_text})

        finally:
            os.remove(file_path)

    return jsonify({"error": "File type not allowed."}), 400

@app.route("/api/delete-document", methods=["POST"])
def delete_document():
    # Get parameters from the request JSON
    data = request.json
    filename = data.get("filename")
    filename = secure_filename(filename)
    tag_name = data.get("tag_name")
    username = data.get("username")

    if not filename and not tag_name:
        return jsonify({"error": "Either 'filename' or 'tag_name' must be provided."}), 400
    
    if not username:
        return jsonify({"error": "username must be provided"}), 400

    # Prepare the SQL condition
    condition = []
    params = []
    
    if filename:
        condition.append("filename = ?")
        params.append(filename)
    
    if tag_name:
        condition.append("tag_name = ?")
        params.append(tag_name)

    if username:
        condition.append("username = ?")
        params.append(username)

    # Combine conditions with 'OR' if both parameters are provided
    condition_str = " AND ".join(condition)

    try:
        # Execute the DELETE query
        session.sql(
            f"DELETE FROM DOCUMENT_CHUNKS WHERE {condition_str}",
            params=params
        ).collect()

        return jsonify({"message": "Document(s) deleted successfully."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=8080)
